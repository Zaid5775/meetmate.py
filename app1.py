# # from flask import Flask, request, jsonify
# # from sumy.parsers.plaintext import PlaintextParser
# # from sumy.nlp.tokenizers import Tokenizer
# # from sumy.summarizers.lex_rank import LexRankSummarizer
# # from flask_cors import CORS

# # app = Flask(__name__)
# # CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
    


# # @app.route("/summarize", methods=["GET", "POST"])
# # def summarize():
# #     if request.method != 'POST':
# #         return "Method not allowed",405
# #     else:
# #         file = request.files['file']
# #         text = file.read().decode('utf-8')
# #           # Create a parser for the text
# #         parser = PlaintextParser.from_string(text, Tokenizer("english"))
    
# #     # Create a summarizer using the LexRank algorithm
# #         summarizer = LexRankSummarizer()
    
# #     # Summarize the text
# #         summary_text = summarizer(parser.document, 5) # 5 is number of sentences to extract
    
# #     # Convert the summary to a string
# #         summary = ' '.join([str(sentence) for sentence in summary_text])
# #         return jsonify(summary)
# # if __name__ == "__main__":
# #     app.run(debug=True)







# from flask import Flask, request, jsonify
# from sumy.parsers.plaintext import PlaintextParser
# from sumy.nlp.tokenizers import Tokenizer
# from sumy.summarizers.lex_rank import LexRankSummarizer
# from flask_cors import CORS
# from docx import Document

# app = Flask(__name__)
# CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
    
# def extract_text_from_docx(file):
#     doc = Document(file)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)

# @app.route("/summarize", methods=["GET", "POST"])
# def summarize():
#     if request.method != 'POST':
#         return "Method not allowed",405
#     else:
#         file = request.files['file']
#         filename = file.filename
#         if filename.endswith('.txt'):
#             text = file.read().decode('utf-8')
#         elif filename.endswith('.docx'):
#             text = extract_text_from_docx(file)
#         else:
#             return "Unsupported file format",400
#           # Create a parser for the text
#         parser = PlaintextParser.from_string(text, Tokenizer("english"))
    
#     # Create a summarizer using the LexRank algorithm
#         summarizer = LexRankSummarizer()
    
#     # Summarize the text
#         summary_text = summarizer(parser.document, 6) # 6 is number of sentences to extract
    
#     # Convert the summary to a string
#         summary = ' '.join([str(sentence) for sentence in summary_text])
#         return jsonify(summary)
# if __name__ == "__main__":
#     app.run(debug=True)





# def prepare_corpus(doc_clean):
#     """
#     Input  : clean document
#     Purpose: create term dictionary of our courpus and Converting list of documents (corpus) into Document Term Matrix
#     Output : term dictionary and Document Term Matrix
#     """
#     # Creating the term dictionary of our courpus, where every unique term is assigned an index. 
#     dictionary = gensim.corpora.Dictionary(doc_clean)
    
#     # Converting list of documents (corpus) into Document Term Matrix using dictionary prepared above.
#     doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
    
#     # generate LDA model
#     ldamodel = gensim.models.ldamodel.LdaModel(doc_term_matrix, num_topics=3, id2word = dictionary, passes=50)
    
#     return ldamodel

# def preprocess(text):
#     result = []
#     for token in gensim.utils.simple_preprocess(text):
#         if token not in gensim.parsing.preprocessing.STOPWORDS and len(token) > 3:
#             result.append(token)
#     return result

# def train_lda(text_data):
#     processed_docs = [preprocess(doc) for doc in text_data]
#     dictionary = gensim.corpora.Dictionary(processed_docs)
#     bow_corpus = [dictionary.doc2bow(doc) for doc in processed_docs]
#     lda_model = gensim.models.LdaMulticore(bow_corpus, num_topics=10, id2word=dictionary, passes=2, workers=2)
#     return lda_model, dictionary, bow_corpus

# def get_document_topics(lda_model, dictionary, bow_corpus, document_index):
#     document_topics = lda_model.get_document_topics(bow_corpus[document_index])
#     topics = []
#     for topic in document_topics:
#         topics.append({'topic_id': topic[0], 'probability': topic[1]})
#     return topics

# def prepare_corpus(doc_clean):
#     # Creating the term dictionary of our courpus, where every unique term is assigned an index. 
#     dictionary = corpora.Dictionary(doc_clean)
    
#     # Converting list of documents (corpus) into Document Term Matrix using dictionary prepared above.
#     doc_term_matrix = [dictionary.doc2bow(doc) for doc in doc_clean]
    
#     # Creating the object for LDA model using gensim library
#     Lda = gensim.models.ldamodel.LdaModel
    
#     # Running and Trainign LDA model on the document term matrix.
#     ldamodel = Lda(doc_term_matrix, num_topics=3, id2word = dictionary, passes=50)
#     return ldamodel



from flask import Flask, request, jsonify
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from flask_cors import CORS
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
import random
import gensim
from gensim.utils import simple_preprocess
from gensim.parsing.preprocessing import STOPWORDS
from gensim import corpora, models
# from transformers import pipeline , set_seed
# from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

# tokenizer = AutoTokenizer.from_pretrained("knkarthick/MEETING_SUMMARY")

# model = AutoModelForSeq2SeqLM.from_pretrained("knkarthick/MEETING_SUMMARY")
from transformers import AutoTokenizer
from transformers import AutoModelForSeq2SeqLM

tokenizer = AutoTokenizer.from_pretrained("knkarthick/MEETING-SUMMARY-BART-LARGE-XSUM-SAMSUM-DIALOGSUM-AMI")


app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})
    

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


@app.route("/summarize", methods=["GET", "POST"])
def summarize():
    if request.method != 'POST':
        return "Method not allowed",405
    else:
        file = request.files['file']
        filename = file.filename
        if filename.endswith('.txt'):
            text = file.read().decode('utf-8')
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file)
        else:
            return "Unsupported file format",400
          # Create a parser for the text
          
        # nltk.download("punkt") # first-time use only
        # doc_clean = [word_tokenize(sent.lower()) for sent in sent_tokenize(text)]  
        # ldamodel = prepare_corpus(doc_clean)
        # topics = ldamodel.print_topics(num_topics=3, num_words=3)
    #     parser = PlaintextParser.from_string(text, Tokenizer("english"))
    
    # # Create a summarizer using the LexRank algorithm
    #     summarizer = LexRankSummarizer()
    # Get the summary length from the request
        summary_length = request.form['summary_length']
        
        minn = 0
        maxx=0
        if summary_length == 'short':
            # num_sentences = random.randint(5,7)
            minn = 1;
            maxx = 1;
        elif summary_length == 'medium':
            # num_sentences = random.randint(7,11)
            minn = 2;
            maxx = 2;
        elif summary_length == 'long':
            # num_sentences = random.randint(11,14)
            minn = 3;
            maxx = 3;
        else:
            # num_sentences = random.randint(7,11)
            minn = 3;
            maxx = 3;
    # Summarize the text
        # summary_text = summarizer(parser.document, num_sentences) # 6 is number of sentences to extract
        # summary_text(text, max_length = maxx, min_length = minn ,do_sample = False)
        inputs = tokenizer(text, return_tensors="pt").input_ids
        model = AutoModelForSeq2SeqLM.from_pretrained("knkarthick/MEETING-SUMMARY-BART-LARGE-XSUM-SAMSUM-DIALOGSUM-AMI")
        outputs = model.generate(inputs, max_length = 100, min_length = 70, do_sample=False)
        summary = tokenizer.decode(outputs[0], skip_special_tokens=True)
        # summary_generator= pipeline('summarization', model="knkarthick/MEETING_SUMMARY")
        # summary = summary_generator(text, max_length=50, min_length=30)[0]["summary_text"]
        # , max_length=maxx, min_length=minn, do_sample=False
        # summary = summary_result['summary']
    
    # Convert the summary to a string
        bullet = '\n- '.join([str(sentence) for sentence in summary])
        bullet_points =  f"- {bullet}"
        # summary = ' '.join([str(sentence) for sentence in summary_text])
        
        original_sentences = sent_tokenize(text)
        original_word_count = 0
        for sentence in original_sentences:
            original_word_count += len(word_tokenize(sentence))
        original_char_count = len(text)
        
        summarized_sentences = sent_tokenize(summary)
        summarized_word_count = 0
        for sentence in summarized_sentences:
            summarized_word_count += len(word_tokenize(sentence))
        summarized_char_count = len(summary)
        
        # nltk.download("punkt") # first-time use only
        # doc_clean = [word_tokenize(sent.lower()) for sent in sent_tokenize(text)]  
        # ldamodel = prepare_corpus(doc_clean)
        # topics = ldamodel.print_topics(num_topics=3, num_words=3)
        
        # //////////////////////
        # text_data = [text]
        # lda_model, dictionary, bow_corpus = train_lda(text_data)
        # document_topics = get_document_topics(lda_model, dictionary, bow_corpus, 0)
        
        return jsonify({
            'summary': summary,
            'bullet_points' : bullet_points,
            'original_sentence_count': len(original_sentences),
            'original_word_count': original_word_count,
            'original_char_count': original_char_count,
            'summarized_sentence_count': len(summarized_sentences),
            'summarized_word_count': summarized_word_count,
            'summarized_char_count': summarized_char_count,
            # 'topics' : topics
            # 'topics' : document_topics
        })
        
if __name__ == "__main__":
    app.run(debug=True)
